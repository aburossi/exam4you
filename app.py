import streamlit as st
import time
import openai
import json
from PyPDF2 import PdfReader
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io
import base64
from io import BytesIO

# Set Streamlit page configuration
st.set_page_config(page_title="Exam Creator", page_icon="📝", layout="wide")

__version__ = "1.3.0"

# --------------------------- Helper Functions ---------------------------

def stream_llm_response(messages, model_params):
    """
    Sends messages to OpenAI's ChatCompletion API and returns the response content.
    """
    openai.api_key = st.secrets["OPENAI_API_KEY"]
    response = openai.ChatCompletion.create(
        model=model_params.get("model", "gpt-4o-mini"),
        messages=messages,
        temperature=model_params.get("temperature", 0.5),
        max_tokens=13000,
    )
    return response.choices[0].message['content']

def extract_text_from_pdf(pdf_file):
    """
    Extracts text content from an uploaded PDF file.
    """
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        extracted_text = page.extract_text()
        if extracted_text:
            text += extracted_text + "\n"
    return text

def extract_text_from_docx(file):
    """
    Extracts text from a DOCX file.
    """
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def process_image(file):
    """
    Processes an uploaded image, resizes it if necessary, and converts it to a Base64-encoded string.
    """
    try:
        image = Image.open(file)
        if image.mode != "RGB":
            image = image.convert("RGB")

        max_size = 1000
        if max(image.size) > max_size:
            image.thumbnail((max_size, max_size))

        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        image_bytes = buffered.getvalue()
        return base64.b64encode(image_bytes).decode("utf-8")
    except Exception as e:
        st.error(f"Error processing image: {e}")
        return None

def chunk_text(text, max_tokens=3000):
    """
    Splits the extracted text into manageable chunks based on the maximum token limit.
    """
    sentences = text.split('. ')
    chunks = []
    chunk = ""
    for sentence in sentences:
        if len(chunk) + len(sentence) > max_tokens:
            chunks.append(chunk)
            chunk = sentence + ". "
        else:
            chunk += sentence + ". "
    if chunk:
        chunks.append(chunk)
    return chunks

def generate_mc_questions(content_text):
    """
    Generates multiple-choice questions based on the provided content using OpenAI's API.
    """
    system_prompt = (
        "Sie sind ein Lehrer für Allgemeinbildung und sollen eine Prüfung zum Thema des eingereichten PDFs erstellen. "
        "Verwenden Sie den Inhalt des PDFs (bitte gründlich analysieren) und erstellen Sie eine Single-Choice-Prüfung auf Oberstufenniveau. "
        "Jede Frage soll genau eine richtige Antwort haben. "
        "Erstellen Sie so viele Prüfungsfragen, wie nötig sind, um den gesamten Inhalt abzudecken, aber maximal 20 Fragen. "
        "Geben Sie die Ausgabe im JSON-Format an. "
        "Das JSON sollte folgende Struktur haben: [{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
        "Stellen Sie sicher, dass das JSON gültig und korrekt formatiert ist."
    )
    user_prompt = (
        "Using the following content from the uploaded PDF, create single-choice questions. "
        "Ensure that each question is based on the information provided in the PDF content and has exactly one correct answer. "
        "Create as many questions as necessary to cover the entire content, but no more than 20 questions. "
        "Provide the output in JSON format with the following structure: "
        "[{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
        "Ensure the JSON is valid and properly formatted.\n\nPDF Content:\n\n"
    ) + content_text

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    try:
        response = stream_llm_response(messages, model_params={"model": "gpt-4o-mini", "temperature": 0.5})
        return response, None
    except Exception as e:
        return None, str(e)

def parse_generated_questions(response):
    """
    Parses the JSON response from OpenAI into Python objects.
    """
    try:
        json_start = response.find('[')
        json_end = response.rfind(']') + 1
        if json_start == -1 or json_end == 0:
            return None, f"No JSON data found in the response. First 500 characters of response:\n{response[:500]}..."
        json_str = response[json_start:json_end]

        questions = json.loads(json_str)
        return questions, None
    except json.JSONDecodeError as e:
        return None, f"JSON parsing error: {e}\n\nFirst 500 characters of response:\n{response[:500]}..."
    except Exception as e:
        return None, f"Unexpected error: {str(e)}\n\nFirst 500 characters of response:\n{response[:500]}..."

# --------------------------- PDF Generation Class ---------------------------

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.cell(0, 10, 'Generated Exam', 0, 1, 'C')

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.multi_cell(0, 10, title)
        self.ln(2)

    def chapter_body(self, body):
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, body)
        self.ln()

    def print_checkbox(self, x, y, checked=False):
        """
        Draws a checkbox at the specified (x, y) coordinates.
        If `checked` is True, the checkbox will be marked.
        """
        # Draw the square for the checkbox
        self.rect(x, y, 5, 5)
        
        if checked:
            # Draw a check mark
            self.set_line_width(0.5)
            self.line(x, y, x + 5, y + 5)
            self.line(x + 5, y, x, y + 5)
            self.set_line_width(0.2)  # Reset to default

# --------------------------- PDF Generation Function ---------------------------

def generate_pdf(questions, include_answers=True):
    """
    Generates a PDF file containing the exam questions.
    """
    pdf = PDF()
    pdf.add_page()

    for i, q in enumerate(questions):
        question = f"Q{i+1}: {q['question']}"
        pdf.chapter_title(question)

        # List the choices
        for choice in q['choices']:
            pdf.chapter_body(choice)

        if include_answers:
            # Add correct answer
            correct_answer = f"Correct answer: {q['correct_answer']}"
            pdf.chapter_body(correct_answer)

            # Add explanation
            explanation = f"Explanation: {q['explanation']}"
            pdf.chapter_body(explanation)

            # Print checkbox for "Test on paper"
            current_y = pdf.get_y()  # Get current y position
            pdf.print_checkbox(10, current_y, True)  # Draw a checked checkbox
            pdf.set_xy(16, current_y)  # Move to the right of the checkbox
            pdf.cell(0, 5, "Test on paper")
            pdf.ln()

    return pdf.output(dest="S").encode("latin1")

# --------------------------- DOCX Generation Function ---------------------------

def generate_docx(questions, include_answers=True):
    """
    Generates a DOCX file containing the exam questions.
    """
    document = Document()
    
    # Set document title
    title = document.add_heading('Generated Exam', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, q in enumerate(questions):
        # Add question number and text
        question_text = f"Q{i+1}: {q['question']}"
        p = document.add_paragraph(question_text, style='List Number')
        
        # Add choices
        for choice in q['choices']:
            p = document.add_paragraph(choice, style='List Bullet')
            p.paragraph_format.left_indent = Pt(20)
        
        if include_answers:
            # Add correct answer
            p = document.add_paragraph("Correct Answer: " + q['correct_answer'])
            p.runs[0].bold = True
            
            # Add explanation
            p = document.add_paragraph("Explanation: " + q['explanation'])
            p.runs[0].italic = True
            
            # Add "Test on paper" checkbox
            p = document.add_paragraph("[ ] Test on paper")
        
        # Add a horizontal line for separation
        document.add_paragraph().add_run().add_break()

    # Save the document to a BytesIO object
    docx_io = BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    
    return docx_io.getvalue()

# --------------------------- Quiz Interaction Functions ---------------------------

def submit_answer(i, quiz_data):
    """
    Handles the submission of an answer for a given question.
    """
    user_choice = st.session_state.get(f"user_choice_{i}")
    
    st.session_state.answers[i] = user_choice

    if user_choice == quiz_data['correct_answer']:
        st.session_state.feedback[i] = ("Correct", quiz_data.get('explanation', 'No explanation available.'))
        st.session_state.correct_answers += 1
    else:
        st.session_state.feedback[i] = ("Incorrect", quiz_data.get('explanation', 'No explanation available.'), quiz_data['correct_answer'])

def mc_quiz_app():
    """
    Renders the multiple-choice quiz interface within the Streamlit app.
    """
    st.subheader('Multiple-Choice Quiz')
    st.write('Please select an answer for each question.')

    questions = st.session_state.generated_questions

    if questions:
        if 'answers' not in st.session_state:
            st.session_state.answers = [None] * len(questions)
            st.session_state.feedback = [None] * len(questions)
            st.session_state.correct_answers = 0

        for i, quiz_data in enumerate(questions):
            st.markdown(f"### Question {i+1}: {quiz_data['question']}")

            if st.session_state.answers[i] is None:
                user_choice = st.radio("Select the correct answer:", quiz_data['choices'], key=f"user_choice_{i}")
                st.button(f"Check Answer {i+1}", key=f"submit_{i}", on_click=submit_answer, args=(i, quiz_data))
            else:
                st.radio("Your answer:", quiz_data['choices'], key=f"user_choice_{i}", index=quiz_data['choices'].index(st.session_state.answers[i]), disabled=True)
                
                feedback_type = st.session_state.feedback[i][0]
                if feedback_type == "Correct":
                    st.success("Correct!")
                else:
                    st.error(f"Incorrect. The correct answer is: {st.session_state.feedback[i][2]}")
                
                st.markdown(f"**Explanation:** {st.session_state.feedback[i][1]}")

        if all(answer is not None for answer in st.session_state.answers):
            score = st.session_state.correct_answers
            total_questions = len(questions)
            st.write(f"""
                <div style="display: flex; flex-direction: column; align-items: center; justify-content: center; height: 100vh;">
                    <h1 style="font-size: 3em; color: gold;">🏆</h1>
                    <h1>Your Score: {score}/{total_questions}</h1>
                </div>
            """, unsafe_allow_html=True)

# --------------------------- Download Functions ---------------------------

def download_files_app():
    """
    Provides options to download the generated quiz as either PDF or DOCX.
    """
    st.subheader('Download Exam as PDF or DOCX')
    
    questions = st.session_state.generated_questions

    if questions:
        # Preview of questions
        with st.expander("Preview Generated Questions"):
            for i, q in enumerate(questions):
                st.markdown(f"### Question {i+1}: {q['question']}")
                for choice in q['choices']:
                    st.write(choice)
                if 'correct_answer' in q:
                    st.write(f"**Correct Answer:** {q['correct_answer']}")
                if 'explanation' in q:
                    st.write(f"**Explanation:** {q['explanation']}")
                st.write("---")
    
        # Choose format and inclusion of answers
        format_option = st.radio("Select the download format:", ["PDF", "DOCX"])
        include_answers = st.checkbox("Include Answers and Explanations", value=True)
        
        if st.button("Generate and Download"):
            if format_option == "PDF":
                file_bytes = generate_pdf(questions, include_answers=include_answers)
                file_name = "exam_with_answers.pdf" if include_answers else "exam_without_answers.pdf"
                mime_type = "application/pdf"
            else:
                file_bytes = generate_docx(questions, include_answers=include_answers)
                file_name = "exam_with_answers.docx" if include_answers else "exam_without_answers.docx"
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            st.download_button(
                label=f"Download {format_option}",
                data=file_bytes,
                file_name=file_name,
                mime=mime_type
            )
    else:
        st.error("No questions found. Please generate an exam first.")

# --------------------------- System Prompt ---------------------------

system_prompt = (
    "Sie sind ein Lehrer für Allgemeinbildung und sollen eine Prüfung zum Thema des eingereichten Inhalts erstellen. "
    "Verwenden Sie den Inhalt (bitte gründlich analysieren) und erstellen Sie eine Single-Choice-Prüfung auf Oberstufenniveau. "
    "Jede Frage soll genau eine richtige Antwort haben. "
    "Erstellen Sie so viele Prüfungsfragen, wie nötig sind, um den gesamten Inhalt abzudecken, aber maximal 20 Fragen. "
    "Geben Sie die Ausgabe im JSON-Format an. "
    "Das JSON sollte folgende Struktur haben: [{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
    "Stellen Sie sicher, dass das JSON gültig und korrekt formatiert ist."
)

# --------------------------- Question Generation Functions ---------------------------

def generate_mc_questions(content, system_prompt):
    """
    Generates multiple-choice questions based on the provided content using OpenAI's API.
    """
    user_prompt = (
        "Using the following content from the uploaded document, create single-choice questions. "
        "Ensure that each question is based on the information provided in the document content and has exactly one correct answer. "
        "Create as many questions as necessary to cover the entire content, but no more than 20 questions. "
        "Provide the output in JSON format with the following structure: "
        "[{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
        "Ensure the JSON is valid and properly formatted.\n\nDocument Content:\n\n" + content
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=1500,
            temperature=0.7,
        )
        return response.choices[0].message["content"], None
    except Exception as e:
        return None, f"Error generating questions: {e}"

def get_questions_from_image(image, system_prompt, user_prompt):
    """
    Generates questions based on an image using OpenAI's API.
    """
    try:
        base64_image = process_image(image)
        if not base64_image:
            return None, "Image processing failed."

        image_payload = {
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{base64_image}",
                "detail": "low",
            },
        }

        messages = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": [{"type": "text", "text": user_prompt}, image_payload],
            },
        ]

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=1500,
            temperature=0.7,
        )
        return response.choices[0].message["content"], None
    except Exception as e:
        return None, f"Error during question generation: {e}"

# --------------------------- File Upload and Question Generation ---------------------------

def pdf_upload_app():
    """
    Handles file upload, content extraction, and question generation.
    """
    st.subheader("Upload Your File - Create Your Exam")
    st.write("Upload a PDF, DOCX, or Image, and we'll handle the rest.")

    uploaded_file = st.file_uploader("Upload a file", type=["pdf", "docx", "jpg", "jpeg", "png"])
    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            with st.spinner("Extracting text from PDF..."):
                content = extract_text_from_pdf(uploaded_file)
            file_type = "text"
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with st.spinner("Extracting text from DOCX..."):
                content = extract_text_from_docx(uploaded_file)
            file_type = "text"
        elif uploaded_file.type.startswith("image/"):
            with st.spinner("Processing image..."):
                content = process_image(uploaded_file)
            file_type = "image"
        else:
            st.error("Unsupported file type.")
            return

        if file_type == "text" and content:
            st.success("File content successfully extracted.")
            st.text_area("Extracted Text (Preview):", value=content[:500] + "...", height=200)

            st.info("Generating exam questions from the uploaded content. This may take a minute...")
            response, error = generate_mc_questions(content, system_prompt)
            if error:
                st.error(error)
            else:
                parsed_questions, parse_error = parse_generated_questions(response)
                if parse_error:
                    st.error(parse_error)
                    st.text_area("Full Response:", value=response, height=200)
                elif parsed_questions:
                    st.session_state.generated_questions = parsed_questions[:20]  # Limit to 20 questions
                    st.success(f"Exam successfully generated with {len(parsed_questions[:20])} questions!")
                    
                    # Display options to proceed
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Take the Exam"):
                            st.session_state.app_mode = "Take Exam"
                            st.experimental_rerun()
                    with col2:
                        if st.button("Download Exam"):
                            st.session_state.app_mode = "Download Exam"
                            st.experimental_rerun()
                else:
                    st.error("No questions were generated. Please try again.")
        elif file_type == "image" and content:
            st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)

            user_prompt = (
                "Using the content derived from the uploaded image, create single-choice questions. "
                "Ensure that each question is based on the image content and has exactly one correct answer. "
                "Create as many questions as necessary to cover the entire content, but no more than 20 questions. "
                "Provide the output in JSON format with the following structure: "
                "[{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
                "Ensure the JSON is valid and properly formatted."
            )

            st.info("Generating exam questions from the uploaded image. This may take a minute...")
            response, error = get_questions_from_image(uploaded_file, system_prompt, user_prompt)
            if error:
                st.error(error)
            else:
                parsed_questions, parse_error = parse_generated_questions(response)
                if parse_error:
                    st.error(parse_error)
                    st.text_area("Full Response:", value=response, height=200)
                elif parsed_questions:
                    st.session_state.generated_questions = parsed_questions[:20]  # Limit to 20 questions
                    st.success(f"Exam successfully generated with {len(parsed_questions[:20])} questions!")
                    
                    # Display options to proceed
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Take the Exam"):
                            st.session_state.app_mode = "Take Exam"
                            st.experimental_rerun()
                    with col2:
                        if st.button("Download Exam"):
                            st.session_state.app_mode = "Download Exam"
                            st.experimental_rerun()
                else:
                    st.error("No questions were generated. Please try again.")
        else:
            st.error("Could not process the uploaded file.")
    else:
        st.warning("Please upload a file to generate an interactive exam.")

# --------------------------- Main Application ---------------------------

def main():
    """
    The main function that controls the Streamlit app's flow.
    """
    st.title("📝 Exam Creator")
    st.markdown(f"**Version:** {__version__}")

    if "app_mode" not in st.session_state:
        st.session_state.app_mode = "Upload PDF & Generate Questions"

    # Define app modes
    app_mode_options = ["Upload File & Generate Questions", "Take Exam", "Download Exam"]
    
    # Sidebar for navigation
    st.sidebar.title("Navigation")
    st.session_state.app_mode = st.sidebar.selectbox("Select App Mode", app_mode_options, index=app_mode_options.index(st.session_state.app_mode) if st.session_state.app_mode in app_mode_options else 0)

    # Render the selected app mode
    if st.session_state.app_mode == "Upload File & Generate Questions":
        pdf_upload_app()
    elif st.session_state.app_mode == "Take Exam":
        if 'generated_questions' in st.session_state and st.session_state.generated_questions:
            mc_quiz_app()
        else:
            st.warning("No generated questions found. Please upload a file and generate questions first.")
    elif st.session_state.app_mode == "Download Exam":
        if 'generated_questions' in st.session_state and st.session_state.generated_questions:
            download_files_app()
        else:
            st.warning("No generated questions found. Please upload a file and generate questions first.")

if __name__ == '__main__':
    main()
